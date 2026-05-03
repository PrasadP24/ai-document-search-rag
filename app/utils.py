import os
from typing import Iterable

from fastapi import HTTPException
from langchain_core.documents import Document
from pypdf import PdfReader


SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".csv", ".xlsx", ".xls"}


def extract_documents(file_path: str, source_name: str) -> list[Document]:
    extension = os.path.splitext(source_name)[1].lower()

    if extension not in SUPPORTED_EXTENSIONS:
        supported = ", ".join(sorted(SUPPORTED_EXTENSIONS))
        raise HTTPException(
            status_code=400,
            detail=f"Unsupported file type '{extension}'. Supported types: {supported}"
        )

    if extension == ".pdf":
        return _extract_pdf(file_path, source_name)
    if extension == ".docx":
        return _extract_docx(file_path, source_name)
    if extension == ".txt":
        return _extract_txt(file_path, source_name)
    if extension == ".csv":
        return _extract_csv(file_path, source_name)

    return _extract_excel(file_path, source_name)


def extract_text_from_pdf(file_path: str) -> str:
    return "\n\n".join(doc.page_content for doc in _extract_pdf(file_path, os.path.basename(file_path)))


def _extract_pdf(file_path: str, source_name: str) -> list[Document]:
    reader = PdfReader(file_path)
    documents = []

    for page_number, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        if text.strip():
            documents.append(
                Document(
                    page_content=text,
                    metadata={
                        "source": source_name,
                        "file_type": "pdf",
                        "page": page_number,
                    },
                )
            )

    return _ensure_content(documents, source_name)


def _extract_docx(file_path: str, source_name: str) -> list[Document]:
    try:
        from docx import Document as DocxDocument
    except ImportError as exc:
        raise HTTPException(status_code=500, detail="python-docx is required for DOCX uploads.") from exc

    docx_file = DocxDocument(file_path)
    parts = []

    for paragraph in docx_file.paragraphs:
        text = paragraph.text.strip()
        if text:
            style = paragraph.style.name if paragraph.style else ""
            prefix = "# " if style.startswith("Heading") else ""
            parts.append(f"{prefix}{text}")

    for table in docx_file.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return _ensure_content(
        [
            Document(
                page_content="\n".join(parts),
                metadata={"source": source_name, "file_type": "docx"},
            )
        ],
        source_name,
    )


def _extract_txt(file_path: str, source_name: str) -> list[Document]:
    with open(file_path, "r", encoding="utf-8", errors="ignore") as file:
        text = file.read()

    return _ensure_content(
        [
            Document(
                page_content=text,
                metadata={"source": source_name, "file_type": "txt"},
            )
        ],
        source_name,
    )


def _extract_csv(file_path: str, source_name: str) -> list[Document]:
    try:
        import pandas as pd
    except ImportError as exc:
        raise HTTPException(status_code=500, detail="pandas is required for CSV uploads.") from exc

    dataframe = pd.read_csv(file_path)
    return _dataframe_to_documents(dataframe, source_name, "csv")


def _extract_excel(file_path: str, source_name: str) -> list[Document]:
    try:
        import pandas as pd
    except ImportError as exc:
        raise HTTPException(status_code=500, detail="pandas and openpyxl are required for Excel uploads.") from exc

    sheets = pd.read_excel(file_path, sheet_name=None)
    documents = []

    for sheet_name, dataframe in sheets.items():
        documents.extend(_dataframe_to_documents(dataframe, source_name, "excel", sheet_name=sheet_name))

    return _ensure_content(documents, source_name)


def _dataframe_to_documents(dataframe, source_name: str, file_type: str, sheet_name: str | None = None) -> list[Document]:
    dataframe = dataframe.fillna("")
    numeric_data = dataframe.apply(_to_numeric_column)
    documents = []

    summary_lines = [
        f"Columns: {', '.join(str(column) for column in dataframe.columns)}",
        f"Rows: {len(dataframe)}",
    ]
    if sheet_name:
        summary_lines.insert(0, f"Sheet: {sheet_name}")

    numeric_columns = numeric_data.dropna(axis=1, how="all")
    if not numeric_columns.empty:
        totals = numeric_columns.sum(numeric_only=True)
        summary_lines.append("Numeric column totals:")
        summary_lines.extend(f"{column}: {value}" for column, value in totals.items())

    documents.append(
        Document(
            page_content="\n".join(summary_lines),
            metadata={
                "source": source_name,
                "file_type": file_type,
                "sheet": sheet_name,
                "row": "summary",
            },
        )
    )

    for row_number, row in dataframe.iterrows():
        values = [f"{column}: {value}" for column, value in row.items() if str(value).strip()]
        if values:
            documents.append(
                Document(
                    page_content="\n".join(values),
                    metadata={
                        "source": source_name,
                        "file_type": file_type,
                        "sheet": sheet_name,
                        "row": int(row_number) + 1,
                    },
                )
            )

    return documents


def _to_numeric_column(column):
    try:
        import pandas as pd
    except ImportError:
        return column

    return pd.to_numeric(column, errors="coerce")


def _ensure_content(documents: Iterable[Document], source_name: str) -> list[Document]:
    cleaned = [doc for doc in documents if doc.page_content.strip()]
    if not cleaned:
        raise HTTPException(status_code=400, detail=f"No readable text found in {source_name}.")

    return cleaned
